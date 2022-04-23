import docx as docx
import streamlit as st
from iapws import IAPWS97 as WSP
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import math as M
from sympy import *
import docx

doc = docx.Document('base.docx')
# pd.set_option("pd.options.display.float_format", 2)

# pd.set_option('float_format' ' .2f '.format)
# pd.set_option('display.float_format', lambda x: '%.2f' % x)

# pd.options.display.float_format = "{:,.2f}".format


st.write("Выполнено: Бастынец А.К. ФПэ-01-19")
doc.add_paragraph("Выполнено: Бастынец А.К. ФПэ-01-19")
st.write("Github: " + "https://github.com/Bastynets/PGT")
doc.add_paragraph("Github: " + "https://github.com/Bastynets/PGT")
st.write("# Задание 1")
doc.add_paragraph("Задание 1")
st.write(
    """Построить процесс расширения пара в турбине. Определение расходов пара на входе в турбину (G0) и в конденсатор (Gк). Получить зависимость КПД ПТУ от параметра заданного в таблице.""")
doc.add_paragraph(
    """Построить процесс расширения пара в турбине. Определение расходов пара на входе в турбину (G0) и в конденсатор (Gк). Получить зависимость КПД ПТУ от параметра заданного в таблице.""")
st.write("""# """)
doc.add_paragraph(" ")
st.write(" *Исходные данные:* ")
doc.add_paragraph(" Исходные данные: ")

age = st.slider('Укажите максимальную границу t0', min_value=500, max_value=550, step=1)
age = age + 0.001  # Посмотреть надобность

Ne = 212e6
p0 = 12.5e6

t0 = list(np.arange(500, age, 1.0))
T0 = [t + 273.15 for t in t0]

ppp = 2.46e6
T0_min = float(T0[0])
T0_max = float(T0[-1])
tpp = 544
Tpp = tpp + 273.15
pk = 4.2e3
tpv = 233
Tpv = tpv + 273.15
delta_p_0 = 0.05 * p0
delta_p_pp = 0.08 * ppp
delta_p = 0.03 * ppp

z = 8

st.write(""" P0 = """ + str(p0 * 10 ** (-6)) + """ МПа""")
doc.add_paragraph(""" P0 = """ + str(p0 * 10 ** (-6)) + """ МПа""")
st.write(""" t0 = """ + str(T0_max - 273.15) + """ C""")
doc.add_paragraph(""" t0 = """ + str(T0_max - 273.15) + """ C""")
st.write(""" Pпп = """ + str(ppp * 10 ** (-6)) + """ МПа""")
doc.add_paragraph(""" Pпп = """ + str(ppp * 10 ** (-6)) + """ МПа""")
st.write(""" tпп = """ + str(tpp) + """ C """)
doc.add_paragraph(""" tпп = """ + str(tpp) + """ C """)
st.write(""" Pк = """ + str(pk * 10 ** (-3)) + """ кПа """)
doc.add_paragraph(""" Pк = """ + str(pk * 10 ** (-3)) + """ кПа """)
st.write(""" tпв = """ + str(tpv) + """ C """)
doc.add_paragraph(""" tпв = """ + str(tpv) + """ C """)
st.write(""" Nэ = """ + str(Ne * 10 ** (-6)) + """ МВт """)
doc.add_paragraph(""" Nэ = """ + str(Ne * 10 ** (-6)) + """ МВт """)
st.write(""" Z = """ + str(z) + """ шт """)
doc.add_paragraph(""" Z = """ + str(z) + """ шт """)

st.write("""# """)
st.write(" *Решение:* ")
doc.add_paragraph(" Решение: ")


def Calculate_eta_G0_Gk(N_e, p_0, T_0, p_pp, T_pp, p_k, T_pv):
    point_0 = WSP(P=p_0 * 10 ** (-6), T=T_0)
    s_0 = point_0.s
    h_0 = point_0.h
    v_0 = point_0.v
    p_0_ = p_0 - 0.05 * p_0
    point_p_0_ = WSP(P=p_0_ * 10 ** (-6), h=h_0)
    t_0_ = point_p_0_.T - 273.15
    s_0_ = point_p_0_.s
    v_0_ = point_p_0_.v

    p_1t = p_pp + 0.1 * p_pp
    point_1t = WSP(P=p_1t * 10 ** (-6), s=s_0)
    t_1t = point_1t.T - 273.15
    h_1t = point_1t.h
    v_1t = point_1t.v

    point_pp = WSP(P=p_pp * 10 ** (-6), T=T_pp)
    h_pp = point_pp.h
    s_pp = point_pp.s
    v_pp = point_pp.v

    H_0 = h_0 - h_1t
    eta_oi = 0.85
    H_i_cvd = H_0 * eta_oi

    h_1 = h_0 - H_i_cvd
    point_1 = WSP(P=p_1t * 10 ** (-6), h=h_1)
    s_1 = point_1.s
    T_1 = point_1.T
    v_1 = point_1.v

    p_pp_ = p_pp - 0.03 * p_pp
    point_pp_ = WSP(P=p_pp_ * 10 ** (-6), h=h_pp)
    s_pp_ = point_pp_.s
    v_pp_ = point_pp_.v
    point_kt = WSP(P=p_k * 10 ** (-6), s=s_pp)
    T_kt = point_kt.T
    h_kt = point_kt.h
    v_kt = point_kt.v
    s_kt = s_pp
    H_0_csdcnd = h_pp - h_kt
    eta_oi = 0.85
    H_i_csdcnd = H_0_csdcnd * eta_oi
    h_k = h_pp - H_i_csdcnd
    point_k = WSP(P=p_k * 10 ** (-6), h=h_k)
    T_k = point_k.T
    s_k = point_k.s
    v_k = point_k.v
    point_k_v = WSP(P=p_k * 10 ** (-6), x=0)
    h_k_v = point_k_v.h
    s_k_v = point_k_v.s
    eta_oiI = (h_1 - h_0) / (h_1t - h_0)
    p_pv = 1.4 * p_0
    point_pv = WSP(P=p_pv * 10 ** (-6), T=T_pv)
    h_pv = point_pv.h
    s_pv = point_pv.s
    ksi_pp_oo = 1 - (1 - (T_k * (s_pp - s_k_v)) / ((h_0 - h_1t) + (h_pp - h_k_v))) / (
            1 - (T_k * (s_pp - s_pv)) / ((h_0 - h_1t) + (h_pp - h_pv)))
    # T_0_= WSP(P = p_pv*10**(-6),x = 0).T
    T_0_ = 374.2 + 273.15
    T_ = (point_pv.T - point_k.T) / (T_0_ - point_k.T)
    if T_ <= 0.636364:
        ksi1 = -1.53 * T_ ** 2 + 2.1894 * T_ + 0.0048
    elif 0.636364 < T_ <= 0.736364:
        ksi1 = -1.3855 * T_ ** 2 + 2.0774 * T_ + 0.0321
    elif 0.736364 < T_ <= 0.863636:
        ksi1 = -2.6535 * T_ ** 2 + 4.2556 * T_ - 0.8569

    if T_ <= 0.631818:
        ksi2 = -1.7131 * T_ ** 2 + 2.3617 * T_ - 0.0142
    elif 0.631818 < T_ <= 0.718182:
        ksi2 = -2.5821 * T_ ** 2 + 3.689 * T_ - 0.4825
    elif 0.718182 < T_ <= 0.827273:
        ksi2 = -1.9864 * T_ ** 2 + 3.138 * T_ - 0.3626
    elif 0.827273 < T_ <= 0.936364:
        ksi2 = -2.0619 * T_ ** 2 + 3.3818 * T_ - 0.4814

    ksi = (ksi1 + ksi2) / 2
    ksi_r_pp = ksi * ksi_pp_oo
    eta_ir = (H_i_cvd + H_i_csdcnd) / (H_i_cvd + (h_pp - h_k_v)) * 1 / (1 - ksi_r_pp)
    H_i = eta_ir * ((h_0 - h_pv) + (h_pp - h_1))
    eta_m = 0.994
    eta_eg = 0.99

    G_0 = N_e / (H_i * eta_m * eta_eg * (10 ** 3))

    G_k = N_e / ((h_k - h_k_v) * eta_m * eta_eg * (10 ** 3)) * (1 / eta_ir - 1)

    return eta_ir, G_0, G_k


eta, G0, Gk = [], [], []
for t in T0:
    eta_ = Calculate_eta_G0_Gk(N_e=Ne, p_0=p0, T_0=t, p_pp=ppp, T_pp=Tpp, p_k=pk, T_pv=Tpv)
    eta.append(eta_[0])
    G0.append(eta_[1])
    Gk.append(eta_[2])

t0_f = [float(x) - 273.15 for x in T0]
eta_f = [float(x) * 100 for x in eta]

st.write(""" Максимальное КПД = """ + str('{:.4}'.format(float(eta_f[-1]))) + """ %""")
doc.add_paragraph(""" Максимальное КПД = """ + str('{:.4}'.format(float(eta_f[-1]))) + """ %""")
st.write(""" Расход пара на входе в турбину (G0) при макс. КПД = """ + str('{:.5}'.format(float(G0[-1]))) + """ кг/с""")
doc.add_paragraph(
    """ Расход пара на входе в турбину (G0) при макс. КПД = """ + str('{:.5}'.format(float(G0[-1]))) + """ кг/с""")
st.write(
    """ Расход пара на входе в конденсатор (Gк) при макс. КПД = """ + str('{:.5}'.format(float(Gk[-1]))) + """ кг/с""")
doc.add_paragraph(
    """ Расход пара на входе в конденсатор (Gк) при макс. КПД = """ + str('{:.5}'.format(float(Gk[-1]))) + """ кг/с""")
st.write("""# """)
st.write(" Табл. Зависимость КПД от t0 ")
doc.add_paragraph(" Табл. Зависимость КПД от t0 ")
# "H_k":round(H_k*1e-3,2),
# t0_eta = pd.DataFrame({"t0, C":[t0_f],
#                        "eta, %":[round(2,eta_f)],
#                        "G_0, кг/с":[round(2,G0)],
#                        "G_k, кг/с":[round(2,Gk)]
#                        })


t0_eta = pd.DataFrame({"t0, C": (t0_f),
                       "eta, %": (eta_f),
                       "G_0, кг/с": (G0),
                       "G_k, кг/с": (Gk)
                       }
                      )


st.dataframe(t0_eta)

#
#
#
# itog=pd.DataFrame({"t0, C":[t0_f],
#                   "eta, %":round(eta_f, 2),
#                   "G_0, кг/с":round(G0, 2),
#                   "G_k, кг/с":round(Gk, 2)})
# itog
#
# t = doc.add_table(itog.shape[0] + 1, itog.shape[1])
#
# # add the header rows.
# for j in range(itog.shape[-1]):
#     t.cell(0, j).text = itog.columns[j]
#
# # add the rest of the data frame
# for i in range(itog.shape[0]):
#     for j in range(itog.shape[-1]):
#         t.cell(i + 1, j).text = str(itog.values[i, j])

st.write("""# """)

t0__eta = plt.figure()

plt.plot(t0_f, eta_f)
plt.plot(t0_f, eta_f, 'ro')
plt.title("Зависимость КПД от начальной температуры")
plt.xlabel("t0, C")
plt.ylabel("КПД, %")
plt.grid()

plt.savefig('1.png')
st.pyplot(t0__eta)
doc.add_picture('1.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

st.title(""" """)

fighs = plt.figure()

point_0 = WSP(P=p0 * 1e-6, T=T0_max)
p_0_d = p0 - delta_p_0
point_0_d = WSP(P=p_0_d * 1e-6, h=point_0.h)
p_1t = ppp + delta_p_pp
point_1t = WSP(P=p_1t * 10 ** (-6), s=point_0.s)
H_01 = point_0.h - point_1t.h
kpd_oi = 0.85
H_i_cvd = H_01 * kpd_oi
h_1 = point_0.h - H_i_cvd
point_1 = WSP(P=p_1t * 1e-6, h=h_1)
p_1 = point_1.P
point_pp = WSP(P=ppp * 1e-6, T=Tpp)
p_pp_d = ppp - delta_p_pp
point_pp_d = WSP(P=p_pp_d * 1e-6, h=point_pp.h)
point_kt = WSP(P=pk * 1e-6, s=point_pp.s)
H_02 = point_pp.h - point_kt.h
kpd_oi = 0.85
H_i_csd_cnd = H_02 * kpd_oi
h_k = point_pp.h - H_i_csd_cnd
point_k = WSP(P=pk * 1e-6, h=h_k)

s_0 = [point_0.s - 0.05, point_0.s, point_0.s + 0.05]
h_0 = [WSP(P=p0 * 1e-6, s=s_).h for s_ in s_0]
s_1 = [point_0.s - 0.05, point_0.s, point_0.s + 0.18]
h_1 = [WSP(P=p_1t * 1e-6, s=s_).h for s_ in s_1]
s_0_d = [point_0_d.s - 0.05, point_0_d.s, point_0_d.s + 0.05]
h_0_d = h_0
s_pp = [point_pp.s - 0.05, point_pp.s, point_pp.s + 0.05]
h_pp = [WSP(P=ppp * 1e-6, s=s_).h for s_ in s_pp]
s_k = [point_pp.s - 0.05, point_pp.s, point_pp.s + 0.8]
h_k = [WSP(P=pk * 1e-6, s=s_).h for s_ in s_k]
s_pp_d = [point_pp_d.s - 0.05, point_pp_d.s, point_pp_d.s + 0.05]
h_pp_d = h_pp

plt.plot([point_0.s, point_0.s, point_0_d.s, point_1.s], [point_1t.h, point_0.h, point_0.h, point_1.h], '-or')
plt.plot([point_pp.s, point_pp.s, point_pp_d.s, point_k.s], [point_kt.h, point_pp.h, point_pp.h, point_k.h], '-or')
plt.plot(s_0, h_0)
plt.plot(s_1, h_1)
plt.plot(s_0_d, h_0_d)
plt.plot(s_pp, h_pp)
plt.plot(s_k, h_k)
plt.plot(s_pp_d, h_pp_d)

for x, y, ind in zip([point_pp.s, point_k.s], [point_pp.h, point_k.h], ['{пп}', '{к}']):
    plt.text(x - 0.45, y + 40, '$h_' + ind + ' = %.2f $' % y)
for x, y, ind in zip([point_kt.s, point_pp_d.s], [point_kt.h, point_pp_d.h], ['{кт}', '{ппд}']):
    plt.text(x + 0.03, y + 40, '$h_' + ind + ' = %.2f $' % y)

for x, y, ind in zip([point_0.s, point_1.s], [point_0.h, point_1.h], ['{0}', '{1}']):
    plt.text(x - 0.01, y + 120, '$h_' + ind + ' = %.2f $' % y)

for x, y, ind in zip([point_1t.s, point_0_d.s], [point_1t.h, point_0_d.h], ['{1т}', '{0д}']):
    plt.text(x + 0.03, y - 60, '$h_' + ind + ' = %.2f $' % y)

    plt.title("h - s диаграмма")
    plt.xlabel("s, кДж/(кг*С)")
    plt.ylabel("h, кДж/кг")
    plt.grid(True)

plt.savefig('2.png')
st.pyplot(fighs)
doc.add_picture('2.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

st.write("# Задание 2")
doc.add_paragraph("Задание 2")

p_0 = 12.5
T_0 = T0_max
d = 1  # m
n = 60  # Гц
G_0 = 170.2137  # кг/с
H_0 = 100
rho = 0.05
alpha_1 = 14  # град
b_1 = 0.06  # м
Delta = 0.003  # м
b_2 = 0.03  # м
kappa_vs = 0
l_1 = 0.015


def iso_bar(wsp_point, min_s=-0.1, max_s=0.11, step_s=0.011, color='b'):
    if not isinstance(wsp_point, list):
        iso_bar_0_s = np.arange(wsp_point.s + min_s, wsp_point.s + max_s, step_s).tolist()
        iso_bar_0_h = [WSP(P=wsp_point.P, s=i).h for i in iso_bar_0_s]
    else:
        iso_bar_0_s = np.arange(wsp_point[0].s + min_s, wsp_point[1].s + max_s, step_s).tolist()
        iso_bar_0_h = [WSP(P=wsp_point[1].P, s=i).h for i in iso_bar_0_s]
    plt.plot(iso_bar_0_s, iso_bar_0_h, color)


def callculate_optimum(d, p_0, T_0, n, G_0, H_0, rho, l_1, alpha_1, b_1, Delta, b_2, kappa_vs):
    u = M.pi * d * n
    point_0 = WSP(P=p_0, T=T_0)
    H_0s = H_0 * (1 - rho)
    H_0r = H_0 * rho
    h_1t = point_0.h - H_0s
    point_1t = WSP(h=h_1t, s=point_0.s)
    c_1t = (2000 * H_0s) ** 0.5
    M_1t = c_1t / point_1t.w
    mu_1 = 0.982 - 0.005 * (b_1 / l_1)
    F_1 = G_0 * point_1t.v / mu_1 / c_1t
    el_1 = F_1 / M.pi / d / M.sin(M.radians(alpha_1))
    e_opt = 5 * el_1 ** 0.5
    if e_opt > 0.85:
        e_opt = 0.85
    l_1 = el_1 / e_opt
    fi_1 = 0.98 - 0.008 * (b_1 / l_1)
    c_1 = c_1t * fi_1
    alpha_1 = M.degrees(M.asin(mu_1 / fi_1 * M.sin(M.radians(alpha_1))))
    w_1 = (c_1 ** 2 + u ** 2 - 2 * c_1 * u * M.cos(M.radians(alpha_1))) ** 0.5
    betta_1 = M.degrees(M.atan(M.sin(M.radians(alpha_1)) / (M.cos(M.radians(alpha_1)) - u / c_1)))
    Delta_Hs = c_1t ** 2 / 2 * (1 - fi_1 ** 2)
    h_1 = h_1t + Delta_Hs * 1e-3
    point_1 = WSP(P=point_1t.P, h=h_1)
    h_2t = h_1 - H_0r
    point_2t = WSP(h=h_2t, s=point_1.s)
    w_2t = (2 * H_0r * 1e3 + w_1 ** 2) ** 0.5
    l_2 = l_1 + Delta
    mu_2 = 0.965 - 0.01 * (b_2 / l_2)
    M_2t = w_2t / point_2t.w
    F_2 = G_0 * point_2t.v / mu_2 / w_2t
    betta_2 = M.degrees(M.asin(F_2 / (e_opt * M.pi * d * l_2)))
    point_1w = WSP(h=point_1.h + w_1 ** 2 / 2 * 1e-3, s=point_1.s)
    psi = 0.96 - 0.014 * (b_2 / l_2)
    w_2 = psi * w_2t
    c_2 = (w_2 ** 2 + u ** 2 - 2 * u * w_2 * M.cos(M.radians(betta_2))) ** 0.5
    alpha_2 = M.degrees(M.atan(M.sin(M.radians(betta_2)) / (M.cos(M.radians(betta_2)) - u / w_2)))
    if alpha_2 < 0:
        alpha_2 = 180 + alpha_2
    Delta_Hr = w_2t ** 2 / 2 * (1 - psi ** 2)
    h_2 = h_2t + Delta_Hr * 1e-3
    point_2 = WSP(P=point_2t.P, h=h_2)
    Delta_Hvs = c_2 ** 2 / 2
    E_0 = H_0 - kappa_vs * Delta_Hvs
    etta_ol1 = (E_0 * 1e3 - Delta_Hs - Delta_Hr - (1 - kappa_vs) * Delta_Hvs) / (E_0 * 1e3)
    etta_ol2 = (u * (c_1 * M.cos(M.radians(alpha_1)) + c_2 * M.cos(M.radians(alpha_2)))) / (E_0 * 1e3)
    return etta_ol2, alpha_2


a1, a2 = [], []
a = callculate_optimum(d, p_0, T_0, n, G_0, H_0, rho, l_1, alpha_1, b_1, Delta, b_2, kappa_vs)
a1.append(a[0])
a2.append(a[1])

H_0 = [i for i in list(range(90, 110, 1))]
alpha1 = []
eta = []
ucf = []
for i in H_0:
    ucf_1 = M.pi * d * n / (2000 * i) ** 0.5
    ucf.append(ucf_1)

    eta_ol, alpha = callculate_optimum(d, p_0, T_0, n, G_0, i, rho, l_1, alpha_1, b_1, Delta, b_2, kappa_vs)
    eta.append(eta_ol)
    alpha1.append(alpha)

plt.plot(ucf, eta)
ucf_eta = plt.figure()
plt.plot(ucf, eta)
plt.title("Зависимость u/cf от etta ")
plt.xlabel("u/cf")
plt.ylabel("etta")
plt.grid()
plt.savefig('3.png')
st.pyplot(ucf_eta)
doc.add_picture('3.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

H_0 = 91  # при этом значении кпд максимальный
u = M.pi * d * n
point_0 = WSP(P=p_0, T=T_0)
H_0s = H_0 * (1 - rho)
H_0r = H_0 * rho
h_1t = point_0.h - H_0s
point_1t = WSP(h=h_1t, s=point_0.s)
c_1t = (2000 * H_0s) ** 0.5
M_1t = c_1t / point_1t.w
mu_1 = 0.982 - 0.005 * (b_1 / l_1)
F_1 = G_0 * point_1t.v / mu_1 / c_1t
el_1 = F_1 / M.pi / d / M.sin(M.radians(alpha_1))
e_opt = 6 * el_1 ** 0.5
if e_opt > 0.85:
    e_opt = 0.85
l_1 = el_1 / e_opt


def plot_hs_nozzle_t(x_lim, y_lim):
    plt.plot([point_0.s, point_1t.s], [point_0.h, point_1t.h], 'bo-')
    iso_bar(point_0, -0.02, 0.02, 0.001, 'g')
    iso_bar(point_1t, -0.02, 0.02, 0.001, 'r')
    plt.xlim(x_lim)
    plt.ylim(y_lim)


plot_hs_nozzle_t([6.12, 6.22], [3200, 3400])
# plt.title("h - s диаграмма")
# plt.xlabel("s, кДж/(кг*С)")
# plt.ylabel("h, кДж/кг")
# plt.grid()
# st.pyplot(hs)
# for x, y, ind in zip([point_0.s, point_1t.s],[point_0.h, point_1t.h], ['{пп}', '{к}']):
# plt.text(x+0.05, y-58, '$h_' + ind + ' = %.2f $'%y)


if alpha_1 <= 10:
    NozzleBlade = 'C-90-09A'
    t1_ = 0.78
    b1_mod = 6.06
    f1_mod = 3.45
    W1_mod = 0.471
    alpha_inst1 = alpha_1 - 12.5 * (t1_ - 0.75) + 20.2
elif 10 < alpha_1 <= 13:
    NozzleBlade = 'C-90-12A'
    t1_ = 0.78
    b1_mod = 5.25
    f1_mod = 4.09
    W1_mod = 0.575
    alpha_inst1 = alpha_1 - 10 * (t1_ - 0.75) + 21.2
elif 13 < alpha_1 <= 16:
    NozzleBlade = 'C-90-15A'
    t1_ = 0.78
    b1_mod = 5.15
    f1_mod = 3.3
    W1_mod = 0.45
    alpha_inst1 = alpha_1 - 16 * (t1_ - 0.75) + 23.1
else:
    NozzleBlade = 'C-90-18A'
    t1_ = 0.75
    b1_mod = 4.71
    f1_mod = 2.72
    W1_mod = 0.333
    alpha_inst1 = alpha_1 - 17.7 * (t1_ - 0.75) + 24.2

z1 = (M.pi * d) / (b_1 * t1_)
z1 = int(z1)
if z1 % 2 == 0:
    print(f'z1 = {z1}')
else:
    z1 = z1 + 1
    print(f'z1 = {z1}')
t1_ = (M.pi * d) / (b_1 * z1)
Ksi_1_ = (0.021042 * b_1 / l_1 + 0.023345) * 100
k_11 = 7.18977510 * M_1t ** 5 - 26.94497258 * M_1t ** 4 + 39.35681781 * M_1t ** 3 - 26.09044664 * M_1t ** 2 + 6.75424811 * M_1t + 0.69896998
k_12 = 0.00014166 * 90 ** 2 - 0.03022881 * 90 + 2.61549380
k_13 = 13.25474043 * t1_ ** 2 - 20.75439502 * t1_ + 9.12762245
Ksi_1 = Ksi_1_ * k_11 * k_12 * k_13

fi_1 = M.sqrt(1 - Ksi_1 / 100)
alpha_1 = 11.8
c_1 = c_1t * fi_1
alpha_1 = M.degrees(M.asin(mu_1 / fi_1 * M.sin(M.radians(alpha_1))))
w_1 = (c_1 ** 2 + u ** 2 - 2 * c_1 * u * M.cos(M.radians(alpha_1))) ** 0.5

c_1u = c_1 * M.cos(M.radians(alpha_1))
c_1a = c_1 * M.sin(M.radians(alpha_1))
w_1u = c_1u - u
w_1_tr = [0, 0, -w_1u, -c_1a]
c_1_tr = [0, 0, -c_1u, -c_1a]
u_1_tr = [-w_1u, -c_1a, -u, 0]
ax = plt.axes()
ax.arrow(*c_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='m', ec='m')
ax.arrow(*w_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='r', ec='r')
ax.arrow(*u_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='y', ec='y')
plt.text(-2 * c_1u / 3, -3 * c_1a / 4, '$c_1$', fontsize=20)
plt.text(-2 * w_1u / 3, -3 * c_1a / 4, '$w_1$', fontsize=20)
plt.title("Cкорости")
plt.xlabel("")
plt.ylabel("")
plt.grid()

st.write(" ")
betta_1 = M.degrees(M.atan(M.sin(M.radians(alpha_1)) / (M.cos(M.radians(alpha_1)) - u / c_1)))

Delta_Hs = c_1t ** 2 / 2 * (1 - fi_1 ** 2)

st.write(" ")
h_1 = h_1t + Delta_Hs * 1e-3
point_1 = WSP(P=point_1t.P, h=h_1)
h_2t = h_1 - H_0r
point_2t = WSP(h=h_2t, s=point_1.s)
w_2t = (2 * H_0r * 1e3 + w_1 ** 2) ** 0.5
l_2 = l_1 + Delta
mu_2 = 0.965 - 0.01 * (b_2 / l_2)
M_2t = w_2t / point_2t.w
F_2 = G_0 * point_2t.v / mu_2 / w_2t
betta_2 = M.degrees(M.asin(F_2 / (e_opt * M.pi * d * l_2)))
point_1w = WSP(h=point_1.h + w_1 ** 2 / 2 * 1e-3, s=point_1.s)

st.write(" ")
if betta_2 <= 15:
    RotorBlade = 'P-23-14A'
    t2_ = 0.63
    b2_mod = 2.59
    f2_mod = 2.44
    W2_mod = 0.39
    beta_inst2 = betta_2 - 12.5 * (t2_ - 0.75) + 20.2

elif 15 < betta_2 <= 19:
    RotorBlade = 'P-26-17A'
    t2_ = 0.65
    b2_mod = 2.57
    f2_mod = 2.07
    W2_mod = 0.225
    beta_inst2 = betta_2 - 19.3 * (t2_ - 0.6) + 60


elif 19 < betta_2 <= 23:
    RotorBlade = 'P-30-21A'
    t2_ = 0.63
    b2_mod = 2.56
    f2_mod = 1.85
    W2_mod = 0.234
    beta_inst2 = betta_2 - 12.8 * (t2_ - 0.65) + 58


elif 23 < betta_2 <= 27:
    RotorBlade = 'P-35-25A'
    t2_ = 0.6
    b2_mod = 2.54
    f2_mod = 1.62
    W2_mod = 0.168
    beta_inst2 = betta_2 - 16.6 * (t2_ - 0.65) + 54.3

elif 27 < betta_2 <= 31:
    RotorBlade = 'P-46-29A'
    t2_ = 0.51
    b2_mod = 2.56
    f2_mod = 1.22
    W2_mod = 0.112
    beta_inst2 = betta_2 - 50.5 * (t2_ - 0.6) + 47.1

else:
    RotorBlade = 'P-50-33A'
    t2_ = 0.49
    b2_mod = 2.56
    f2_mod = 1.02
    W2_mod = 0.079
    beta_inst2 = betta_2 - 20.8 * (t2_ - 0.6) + 43.7

z2 = int((M.pi * d) / (b_2 * t2_))

t2_ = (M.pi * d) / (b_2 * z2)
Ksi_2_ = 4.364 * b_2 / l_2 + 4.22
k_21 = -13.79438991 * M_2t ** 4 + 36.69102267 * M_2t ** 3 - 32.78234341 * M_2t ** 2 + 10.61998662 * M_2t + 0.28528786
k_22 = 0.00331504 * betta_1 ** 2 - 0.21323910 * betta_1 + 4.43127194
k_23 = 60.72813684 * t2_ ** 2 - 76.38053189 * t2_ + 24.97876023
Ksi_2 = Ksi_2_ * k_21 * k_22 * k_23

psi = M.sqrt(1 - Ksi_2 / 100)

st.write(" ")
st.write(" ")

psi = 0.94
w_2 = psi * w_2t
c_2 = (w_2 ** 2 + u ** 2 - 2 * u * w_2 * M.cos(M.radians(betta_2))) ** 0.5
alpha_2 = M.degrees(M.atan(M.sin(M.radians(betta_2)) / (M.cos(M.radians(betta_2)) - u / w_2)))
Delta_Hr = w_2t ** 2 / 2 * (1 - psi ** 2)
h_2 = h_2t + Delta_Hr * 1e-3
point_2 = WSP(P=point_2t.P, h=h_2)
Delta_Hvs = c_2 ** 2 / 2
E_0 = H_0 - kappa_vs * Delta_Hvs
etta_ol1 = (E_0 * 1e3 - Delta_Hs - Delta_Hr - (1 - kappa_vs) * Delta_Hvs) / (E_0 * 1e3)
etta_ol2 = (u * (c_1 * M.cos(M.radians(alpha_1)) + c_2 * M.cos(M.radians(alpha_2)))) / (E_0 * 1e3)

h_vs = h_2 + Delta_Hvs * 1e-3
point_vs = WSP(P=point_2t.P, h=h_vs)

hsstage = plt.figure()


def plot_hs_stage_t(x_lim, y_lim):
    plot_hs_nozzle_t(x_lim, y_lim)
    plt.plot([point_0.s, point_1.s], [point_0.h, point_1.h], 'bo-')
    plt.plot([point_1.s, point_2t.s], [point_1.h, point_2t.h], 'ro-')
    iso_bar(point_2t, -0.02, 0.02, 0.001, 'y')
    plt.plot([point_2.s, point_vs.s], [point_2.h, point_vs.h], 'ro-')
    plt.plot([point_1.s, point_2.s], [point_1.h, point_2.h], 'bo-')


plot_hs_stage_t([6.61, 6.66], [3360, 3500])
plt.title("h - s диаграмма")
plt.xlabel("s, кДж/(кг*С)")
plt.ylabel("h, кДж/кг")
plt.grid()
plt.savefig('4.png')
st.pyplot(hsstage)
doc.add_picture('4.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

st.write(" ")
cw = plt.figure()
c_1u = c_1 * M.cos(M.radians(alpha_1))
c_1a = c_1 * M.sin(M.radians(alpha_1))
w_1u = c_1u - u
w_2a = w_2 * M.sin(M.radians(betta_2))
w_2u = w_2 * M.cos(M.radians(betta_2))
c_2u = w_2u + u
w_1_tr = [0, 0, -w_1u, -c_1a]
c_1_tr = [0, 0, -c_1u, -c_1a]
u_1_tr = [-w_1u, -c_1a, -u, 0]

w_2_tr = [0, 0, w_2u, -w_2a]
c_2_tr = [0, 0, c_2u, -w_2a]
u_2_tr = [c_2u, -w_2a, -u, 0]
ax = plt.axes()
ax.arrow(*c_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='r', ec='r')
ax.arrow(*w_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='b', ec='b')
ax.arrow(*u_1_tr, head_width=5, length_includes_head=True, head_length=20, fc='g', ec='g')
ax.arrow(*c_2_tr, head_width=5, length_includes_head=True, head_length=20, fc='r', ec='r')
ax.arrow(*w_2_tr, head_width=5, length_includes_head=True, head_length=20, fc='b', ec='b')
ax.arrow(*u_2_tr, head_width=5, length_includes_head=True, head_length=20, fc='g', ec='g')
plt.text(-2 * c_1u / 3, -3 * c_1a / 4, '$c_1$', fontsize=20)
plt.text(-2 * w_1u / 3, -3 * c_1a / 4, '$w_1$', fontsize=20)
plt.text(2 * c_2u / 3, -3 * w_2a / 4, '$c_2$', fontsize=20)
plt.text(2 * w_2u / 3, -3 * w_2a / 4, '$w_2$', fontsize=20)
plt.title("Треугольник скоростей")
plt.xlabel("")
plt.ylabel("")
plt.grid()
plt.savefig('4.png')
st.pyplot(cw)
doc.add_picture('4.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

st.write(" ")
delta_a = 0.0025
z_per_up = 2
mu_a = 0.5
mu_r = 0.75
d_per = d + l_1
delta_r = d_per * 0.001
delta_ekv = 1 / M.sqrt(1 / (mu_a * delta_a) ** 2 + z_per_up / (mu_r * delta_r) ** 2)

xi_u_b = M.pi * d_per * delta_ekv * etta_ol1 / F_1 * M.sqrt(rho + 1.8 * l_2 / d)

Delta_Hub = xi_u_b * E_0

st.write(" ")
st.write(" ")
st.write(" ")

k_tr = 0.0007
Kappa_VS = 0
u = M.pi * d * n
c_f = M.sqrt(2000 * H_0)
ucf = u / c_f
xi_tr = k_tr * d ** 2 / F_1 * ucf ** 3

Delta_Htr = xi_tr * E_0

st.write(" ")
st.write(" ")
k_v = 0.065
m = 1
xi_v = k_v / M.sin(M.radians(alpha_1)) * (1 - e_opt) / e_opt * ucf ** 3 * m

i_p = 4
B_2 = b_2 * M.sin(M.radians(beta_inst2))
xi_segm = 0.25 * B_2 * l_2 / F_1 * ucf * etta_ol1 * i_p

xi_parc = xi_v + xi_segm
Delta_H_parc = E_0 * xi_parc

H_i = E_0 - Delta_Hr * 1e-3 - Delta_Hs * 1e-3 - (1 - Kappa_VS) * Delta_Hvs * 1e-3 - Delta_Hub - Delta_Htr - Delta_H_parc

eta_oi = H_i / E_0
st.write("""Внутренний относительный КПД ступени eta_oi  = %.3f """ % eta_oi)
doc.add_paragraph("""Внутренний относительный КПД ступени eta_oi  = %.3f """ % eta_oi)
N_i = G_0 * H_i
st.write("""Внутреняя мощность ступени N_i =  %.2f кВт""" % N_i)
doc.add_paragraph("""Внутреняя мощность ступени N_i =  %.2f кВт""" % N_i)

# st.write(point_0_d.P) #точка которая покажет p0 11.875 МПа
# st.write(point_0_d.h) #точка которая покажет h0 3343.566271475289 кдж
# st.write(point_1.P) #давление за цвд 9.464835340804248 Мпа
# st.write(eta_f)
# st.write(G0)


st.write("# Задание 3")
doc.add_paragraph(" Задание 3")
st.write("""Определение числа ступеней и распределение параметров по ним""")
doc.add_paragraph("""Определение числа ступеней и распределение параметров по ним""")
st.write("""# """)

# Исходные данные

drs = 1  # Введите диаметр регулирующей ступени, м
P0 = point_0_d.P  # Введите давление полного торможения перед нерегулируемой ступенью P0, МПа
h0 = point_0_d.h  # Введите энтальпия полного торможения перед первой нерегулируемой ступенью h0, кДж/кг
G0 = G_0  # Введите расход пара в первую нерегулируемую ступень, кг/с
etaoi = eta_oi  # Введите внутренний КПД ЦВД
Z = 8  # Введите количество ступеней ЦВД
Pz = p_1  # Введите давление за ЦВД, МПа

deltaD = 0.26  # m
n = 50  # Гц
rho_s = 0.05
alfa = 15  # град
fi = 0.96
mu1 = 0.97
delta = 0.003
tetta = 20

st.write("""# """)
st.write(" *Дано:* ")
doc.add_paragraph(" Дано:* ")
st.write(""" P0 = """ + str('{:.3}'.format(float(P0))) + """ МПа""")
doc.add_paragraph(""" P0 = """ + str('{:.3}'.format(float(P0))) + """ МПа""")
st.write(""" h0 = """ + str('{:.2f}'.format(float(h0))) + """ кДж/кг""")
doc.add_paragraph(""" h0 = """ + str('{:.2f}'.format(float(h0))) + """ кДж/кг""")
st.write(""" dрс = """ + str(drs) + """ м """)
doc.add_paragraph(""" dрс = """ + str(drs) + """ м """)
st.write(""" Z = """ + str(Z) + """ шт """)
doc.add_paragraph(""" Z = """ + str(Z) + """ шт """)
st.write(""" G0 = """ + str('{:.2f}'.format(float(G0))) + """ кг/с """)
doc.add_paragraph(""" G0 = """ + str('{:.2f}'.format(float(G0))) + """ кг/с """)
st.write(""" n = """ + str(n) + """ Гц """)
doc.add_paragraph(""" n = """ + str(n) + """ Гц """)
st.write(""" eta_oi = """ + str('{:.2}'.format(float(etaoi))) + """ """)
doc.add_paragraph(""" eta_oi = """ + str('{:.2}'.format(float(etaoi))) + """ """)
st.write(""" Pz = """ + str('{:.3}'.format(float(Pz))) + """ МПа """)
doc.add_paragraph(""" Pz = """ + str('{:.3}'.format(float(Pz))) + """ МПа """)

st.write("""# """)
st.write(" *Решение:* ")
doc.add_paragraph("Решение:")

D1 = drs - deltaD
sat_steam = WSP(P=P0, h=h0)
s_0 = sat_steam.s
t_0 = sat_steam.T

error = 2
i = 1
while error > 0.5:
    rho = rho_s + 1.8 / (tetta + 1.8)
    X = (fi * M.cos(M.radians(alfa))) / (2 * M.sqrt(1 - rho))
    H01 = 12.3 * (D1 / X) ** 2 * (n / 50) ** 2
    h2t = h0 - H01
    steam2t = WSP(h=h2t, s=s_0)
    v2t = steam2t.v
    l11 = G0 * v2t * X / (M.pi ** 2 * D1 ** 2 * n * M.sqrt(1 - rho) * M.sin(M.radians(alfa)) * mu1)
    tetta_old = tetta
    tetta = D1 / l11
    # print(i, tetta_old, tetta)
    error = abs(tetta - tetta_old) / tetta_old * 100
    # print(error)
    i += 1

l21 = l11 + delta
d_s = D1 - l21
steam_tz = WSP(P=Pz, s=s_0)
h_zt = steam_tz.h
H0 = h0 - h_zt
Hi = H0 * etaoi
h_z = h0 - Hi
steam_z = WSP(P=Pz, h=h_z)
v_2z = steam_z.v
x = Symbol('x')
с = solve(x ** 2 + x * d_s - (l21 * (d_s + l21) * v_2z / v2t))
for j in с:
    if j > 0:
        l2z = j
d2z = d_s + l2z
tetta1 = (l21 + d_s) / l21
tettaz = (l2z + d_s) / l2z
rho1 = rho_s + 1.8 / (1.8 + tetta1)
rhoz = rho_s + 1.8 / (1.8 + tettaz)
X1 = (fi * cos(M.radians(alfa))) / (2 * sqrt(1 - rho1))
Xz = (fi * cos(M.radians(alfa))) / (2 * sqrt(1 - rhoz))

DeltaZ = 1
ite = 0
while DeltaZ > 0:
    matr = []
    Num = 0
    SumH = 0
    for _ in range(int(Z)):
        li = (l21 - l2z) / (1 - Z) * Num + l21
        di = (D1 - d2z) / (1 - Z) * Num + D1
        tetta_i = di / li
        rho_i = rho_s + 1.8 / (1.8 + tetta_i)
        X_i = (fi * M.cos(M.radians(alfa))) / (2 * M.sqrt(1 - rho_i))
        if Num < 1:
            H_i = 12.3 * (di / X_i) ** 2 * (n / 50) ** 2
        else:
            H_i = 12.3 * (di / X_i) ** 2 * (n / 50) ** 2 * 0.95
        Num = Num + 1
        H_d = 0
        SumH = SumH + H_i
        matr.append([Num, round(di, 3), round(li, 3), round(tetta_i, 2), round(rho_i, 3), round(X_i, 3), round(H_i, 2),
                     round(H_d, 2)])
    H_m = SumH / Z
    q_t = 4.8 * 10 ** (-4) * (1 - etaoi) * H0 * (Z - 1) / Z
    Z_new = round(H0 * (1 + q_t) / H_m)
    DeltaZ = abs(Z - Z_new)
    # print(ite, Z)
    Z = Z_new
    ite += 1
DeltaH = (H0 * (1 + q_t) - SumH) / Z
a = 0
for elem in matr:
    matr[a][7] = round(elem[6] + DeltaH, 2)
    a += 1

## Добавлено для таблицы
N_ = []
di_ = []
li_ = []
tettai_ = []
rhoi_ = []
Xi_ = []
Hi_ = []
Hdi_ = []
a = 0
for elem in matr:
    N_.append(matr[a][0])
    di_.append(matr[a][1])
    li_.append(matr[a][2])
    tettai_.append(matr[a][3])
    rhoi_.append(matr[a][4])
    Xi_.append(matr[a][5])
    Hi_.append(matr[a][6])
    Hdi_.append(matr[a][7])
    a += 1

di_ = [float(x) for x in di_]
li_ = [float(x) for x in li_]
tettai_ = [float(x) for x in tettai_]
rhoi_ = [float(x) for x in rhoi_]
Xi_ = [float(x) for x in Xi_]
Hi_ = [float(x) for x in Hi_]
Hdi_ = [float(x) for x in Hdi_]

## Таблица
# table = pd.DataFrame({"№ ст":(N_),
#                       "di, м":round(di_,3),
#                       "li, м":round(li_,3),
#                       "θi ":round(tettai_,2),
#                       "ρi ":round(rhoi_,3),
#                       "Xi ":round(Xi_),
#                       "Hi, кДж/кг":round(Hi_,2),
#                       "Hi + Δ, кДж/кг":round(Hdi_,2)
#                       }
#                      )
table = pd.DataFrame({"№ ст": (N_),
                      "di, м": (di_),
                      "li, м": (li_),
                      "θi ": (tettai_),
                      "ρi ": (rhoi_),
                      "Xi ": (Xi_),
                      "Hi, кДж/кг": (Hi_),
                      "Hi + Δ, кДж/кг": (Hdi_)
                      }
                     )
st.dataframe(table)

# t_ = doc.add_table(table.shape[0] + 1, table.shape[1])
#
# # add the header rows.
# for j in range(table.shape[-1]):
#     t_.cell(0, j).text = table.columns[j]
#
# # add the rest of the data frame
# for i in range(table.shape[0]):
#     for j in range(table.shape[-1]):
#         t_.cell(i + 1, j).text = str(table.values[i, j])
#
# ## Графики
z = []
for a in range(1, Z + 1):
    z.append(a)

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, di_, '-ro')
plt.title('Рисунок 1. Распределение средних диаметров по проточной части')
plt.savefig('11.png')
st.pyplot(fig)
doc.add_picture('11.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, li_, '-ro')
plt.title('Рисунок 2. Распределение высот лопаток по проточной части')
plt.savefig('12.png')
st.pyplot(fig)
doc.add_picture('12.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, tettai_, '-ro')
plt.title('Рисунок 3. Распределение обратной веерности по проточной части')
plt.savefig('13.png')
st.pyplot(fig)
doc.add_picture('13.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, rhoi_, '-ro')
plt.title('Рисунок 4. Распределение степени реактивности по проточной части')
plt.savefig('14.png')
st.pyplot(fig)
doc.add_picture('14.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, Xi_, '-ro')
plt.title('Рисунок 5. Распределение U/Cф по проточной части')
plt.savefig('15.png')
st.pyplot(fig)
doc.add_picture('15.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, Hi_, '-ro')
plt.title('Рисунок 6. Распределение теплоперепадов по проточной части')
plt.savefig('16.png')
st.pyplot(fig)
doc.add_picture('16.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

st.write("#")
fig = plt.figure(figsize=(10, 5))
ax = fig.gca()
ax.set_xticks(np.arange(1, 20, 1))
plt.grid(True)
plt.plot(z, Hdi_, '-ro')
plt.title('Рисунок 7. Распределение теплоперепадов с учетом невязки по проточной части')
plt.savefig('17.png')
st.pyplot(fig)
doc.add_picture('17.png', width=docx.shared.Inches(6), height=docx.shared.Inches(3))

# ВЫВОД ПАРМЕТРОВ В ВОРД doc.remove_paragraph
# doc = docx.Document('base.docx')
#
# doc.add_paragraph("""Внутренний относительный КПД ступени   eta_oi  = %.3f """ % eta_oi)
#
# doc.add_picture('3.png', width=docx.shared.Inches(4), height=docx.shared.Inches(3))

doc.save('result.docx')

with open("result.docx", "rb") as file:
    st.download_button(
        label="Download Result",
        data=file,
        file_name='result.docx',
    )
